"""
Dobby — Document Builder Service
Handles generating tailored resumes and cover letters using LLMs,
and constructing physical `.docx` files using python-docx.
"""

import json
import logging
from typing import Dict, Any, Optional
import google.generativeai as genai
import httpx
from pydantic import BaseModel, Field
from io import BytesIO

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from config.settings import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()

if settings.gemini_api_key:
    genai.configure(api_key=settings.gemini_api_key)

# ─── Pydantic Schemas for Structured LLM Output ────────────────

class ExperienceItem(BaseModel):
    title: str = Field(description="Job title or project name")
    company: str = Field(description="Company name or 'Independent Project'")
    duration: str = Field(description="Duration (e.g., 'Jan 2020 - Present')")
    bullets: list[str] = Field(description="3-5 impactful, results-oriented bullet points tailored to the target job.")

class EducationItem(BaseModel):
    degree: str
    institution: str
    graduation_year: str

class TailoredResume(BaseModel):
    name: str
    contact_info: str = Field(description="E.g., Email | Phone | Location | LinkedIn")
    summary: str = Field(description="A 2-3 sentence professional summary highlighting fit for the specific role.")
    skills: list[str] = Field(description="A categorized or comma-separated list of top skills relevant to the job.")
    experience: list[ExperienceItem]
    education: list[EducationItem]
    projects: list[ExperienceItem]

class CoverLetterData(BaseModel):
    body: str = Field(description="A modern, punchy 3-4 paragraph cover letter body mapping the candidate's skills to the job description.")


# ─── LLM Generation Functions ────────────────────────────────────

def _build_resume_prompt(profile: Dict[str, Any], job_details: Dict[str, Any]) -> str:
    jd = job_details.get("description", "")[:6000]
    return f"""You are an expert Executive Resume Writer.
Your task is to take the candidate's core profile and output a tailored, highly impactful resume data structure targeting the provided job description.

# Candidate Profile
Name: {profile.get("name")}
Location: {profile.get("location")}
Education: {profile.get("education")}
Experience: {profile.get("experience")}
Skills: {", ".join(profile.get("skills", []))}
Projects/Portfolio: {profile.get("projects")}

# Job Posting
Role/Title: {job_details.get("role")}
Company: {job_details.get("company")}
Description:
{jd}

Re-write the candidate's summary and experience bullets to emphasize achievements that align precisely with the job description. Do NOT hallucinate jobs they did not hold, but frame their existing experience in the best possible light for this specific role. Use action verbs and metric-driven points where possible.
"""

def _build_cover_letter_prompt(profile: Dict[str, Any], job_details: Dict[str, Any]) -> str:
    jd = job_details.get("description", "")[:6000]
    return f"""You are an expert Career Coach writing a compelling cover letter.
Write the body paragraphs of a modern, punchy cover letter for the following candidate applying to the following job.

# Candidate Profile
Name: {profile.get("name")}
Experience: {profile.get("experience")}
Skills: {", ".join(profile.get("skills", []))}

# Job Posting
Role/Title: {job_details.get("role")}
Company: {job_details.get("company")}
Description:
{jd}

Rules:
- Write ONLY the body of the letter (no 'Dear Hiring Manager' or 'Sincerely, Name' — we will add those).
- Max 3-4 paragraphs.
- Be confident, direct, and highlight specific alignment between the candidate's background and the job's core needs.
- Avoid clichés like "I am writing to express my interest in..." Start with a strong hook.
"""

async def generate_tailored_resume(profile: Dict[str, Any], job_details: Dict[str, Any]) -> Dict[str, Any]:
    """Uses Gemini to generate structured tailored resume content."""
    if not settings.gemini_api_key:
        raise ValueError("Gemini API key not configured.")
        
    prompt = _build_resume_prompt(profile, job_details)
    
    model = genai.GenerativeModel("gemini-2.5-flash")
    response = model.generate_content(
        prompt,
        generation_config=genai.GenerationConfig(
            response_mime_type="application/json",
            response_schema=TailoredResume,
            temperature=0.3,
        ),
    )
    return json.loads(response.text)

async def generate_cover_letter(profile: Dict[str, Any], job_details: Dict[str, Any]) -> str:
    """Uses Gemini to generate a tailored cover letter body."""
    if not settings.gemini_api_key:
        raise ValueError("Gemini API key not configured.")
        
    prompt = _build_cover_letter_prompt(profile, job_details)
    
    model = genai.GenerativeModel("gemini-2.5-flash")
    response = model.generate_content(
        prompt,
        generation_config=genai.GenerationConfig(
            response_mime_type="application/json",
            response_schema=CoverLetterData,
            temperature=0.4,
        ),
    )
    data = json.loads(response.text)
    return data.get("body", "")


# ─── DOCX Construction ──────────────────────────────────────────

def create_docx_from_resume_data(data: Dict[str, Any]) -> BytesIO:
    """Builds a minimalist, ATS-friendly MS Word docx file in memory."""
    doc = Document()
    
    # Setup margins (narrow)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Header: Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(data.get("name", "Candidate Name"))
    name_run.bold = True
    name_run.font.size = Pt(18)
    
    # Header: Contact
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.add_run(data.get("contact_info", ""))
    
    doc.add_paragraph() # spacing
    
    # ── Summary ──
    summary = data.get("summary")
    if summary:
        _add_section_heading(doc, "PROFESSIONAL SUMMARY")
        doc.add_paragraph(summary)
    
    # ── Skills ──
    skills = data.get("skills", [])
    if skills:
        _add_section_heading(doc, "SKILLS & EXPERTISE")
        doc.add_paragraph(" • ".join(skills))
        doc.add_paragraph() # spacing
        
    # ── Experience ──
    experience = data.get("experience", [])
    if experience:
        _add_section_heading(doc, "PROFESSIONAL EXPERIENCE")
        for item in experience:
            # Title & Duration
            p1 = doc.add_paragraph()
            p1.paragraph_format.space_after = Pt(2)
            run_title = p1.add_run(item.get("title", ""))
            run_title.bold = True
            p1.add_run(" | " + item.get("company", ""))
            
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p2.add_run(item.get("duration", "")).italic = True
            
            # Bullets
            for bullet in item.get("bullets", []):
                bp = doc.add_paragraph(bullet, style='List Bullet')
                bp.paragraph_format.space_after = Pt(2)
                
            doc.add_paragraph() # spacing

    # ── Projects ──
    projects = data.get("projects", [])
    if projects:
        _add_section_heading(doc, "RELEVANT PROJECTS")
        for item in projects:
            p1 = doc.add_paragraph()
            p1.paragraph_format.space_after = Pt(2)
            p1.add_run(item.get("title", "")).bold = True
            for bullet in item.get("bullets", []):
                bp = doc.add_paragraph(bullet, style='List Bullet')
                bp.paragraph_format.space_after = Pt(2)
            doc.add_paragraph() # spacing

    # ── Education ──
    education = data.get("education", [])
    if education:
        _add_section_heading(doc, "EDUCATION")
        for ed in education:
            doc.add_paragraph(f"{ed.get('degree', '')} — {ed.get('institution', '')} ({ed.get('graduation_year', '')})")

    # Save to BytesIO stream
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def _add_section_heading(doc: Document, text: str):
    """Adds a bold heading with a bottom border (using simple text/style for ATS)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    # Simple line
    doc.add_paragraph("_" * 70).paragraph_format.space_after = Pt(6)
